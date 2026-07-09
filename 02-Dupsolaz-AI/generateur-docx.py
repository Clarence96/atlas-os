from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import json
import os

# ============================================================
# CONFIGURATION
# ============================================================
DOSSIER_SORTIE = os.path.join(os.path.dirname(__file__), "documents_generees")
DOSSIER_ASSETS = os.path.join(os.path.dirname(__file__), "..", "00-Foundation", "Assets")
CHEMIN_LOGO = os.path.join(DOSSIER_ASSETS, "logo-dupsolaz.png")

os.makedirs(DOSSIER_SORTIE, exist_ok=True)

# ============================================================
# FONCTIONS
# ============================================================
def ajouter_logo(doc):
    """Ajoute le logo Dupsolaz en haut du document."""
    if os.path.exists(CHEMIN_LOGO):
        logo_par = doc.add_paragraph()
        logo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_par.add_run()
        run.add_picture(CHEMIN_LOGO, width=Inches(2.0))
        doc.add_paragraph()  # Espacement


def ajouter_entete_dupsolaz(doc):
    """Ajoute l'en-tête avec les infos Dupsolaz."""
    # Titre principal
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre.add_run("DUPSOLAZ LEGACY")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Sous-titre
    sous_titre = doc.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sous_titre.add_run("Services Administratifs & Structuration d'Entreprise")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()


def ajouter_infos_emetteur(doc, emetteur):
    """Ajoute le bloc informations de Dupsolaz."""
    doc.add_paragraph("ÉMETTEUR", style='Heading 3')
    
    infos = doc.add_paragraph()
    infos.add_run(f"{emetteur.get('nom', '')}\n").bold = True
    infos.add_run(f"RIDET : {emetteur.get('ridet', '')}\n")
    infos.add_run(f"{emetteur.get('forme_juridique', '')}\n")
    infos.add_run(f"{emetteur.get('adresse', '')}\n")
    infos.add_run(f"Email : {emetteur.get('email', '')}\n")
    infos.add_run(f"Tél : {emetteur.get('telephone', '')}")
    
    doc.add_paragraph()


def ajouter_infos_client(doc, client):
    """Ajoute le bloc informations client."""
    doc.add_paragraph("CLIENT", style='Heading 3')
    
    infos = doc.add_paragraph()
    infos.add_run(f"{client.get('nom', '')}\n").bold = True
    infos.add_run(f"{client.get('adresse', '')}\n")
    infos.add_run(f"Email : {client.get('email', '')}\n")
    infos.add_run(f"Tél : {client.get('telephone', '')}")
    
    doc.add_paragraph()


def ajouter_tableau_prestations(doc, prestations, total_ht, tgc, total_ttc):
    """Ajoute le tableau des prestations."""
    doc.add_paragraph("PRESTATIONS", style='Heading 3')
    
    # Création du tableau
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # En-têtes
    headers = ['Prestation', 'Quantité', 'Unité', 'Prix unitaire (XPF)', 'Montant (XPF)']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Lignes de prestations
    for prestation in prestations:
        row = table.add_row()
        row.cells[0].text = prestation.get('libelle', '')
        row.cells[1].text = str(prestation.get('quantite', ''))
        row.cells[2].text = prestation.get('unite', '')
        row.cells[3].text = f"{prestation.get('tarif_unitaire', 0):,}"
        row.cells[4].text = f"{prestation.get('montant', 0):,}"
    
    doc.add_paragraph()
    
    # Totaux
    totaux = doc.add_paragraph()
    totaux.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    totaux.add_run(f"Total HT : {total_ht:,} XPF\n").bold = True
    totaux.add_run(f"TGC ({tgc}%) : {total_ttc - total_ht:,} XPF\n")
    totaux.add_run(f"Total TTC : {total_ttc:,} XPF\n").bold = True


def ajouter_zone_signature(doc):
    """Ajoute une zone de signature pour le client et l'émetteur."""
    doc.add_paragraph()
    doc.add_paragraph("SIGNATURES", style='Heading 3')
    
    # Création d'un tableau invisible pour aligner les signatures
    table_sig = doc.add_table(rows=2, cols=2)
    table_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # En-têtes des colonnes
    cell_client = table_sig.rows[0].cells[0]
    cell_client.text = "Le Client"
    for p in cell_client.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
    
    cell_emetteur = table_sig.rows[0].cells[1]
    cell_emetteur.text = "Dupsolaz Legacy"
    for p in cell_emetteur.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
    
    # Zones de signature (lignes vides)
    for col in range(2):
        cell = table_sig.rows[1].cells[col]
        cell.text = "\n\n\n\n\n"  # Espace pour la signature
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Mention sous les signatures
    doc.add_paragraph()
    mention = doc.add_paragraph()
    mention.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mention.add_run("Signature précédée de la mention « Bon pour accord »")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)


def ajouter_mentions(doc, conditions, mentions_legales, date_echeance):
    """Ajoute les conditions et mentions légales."""
    doc.add_paragraph("CONDITIONS", style='Heading 3')
    
    if conditions:
        for condition in conditions:
            doc.add_paragraph(condition, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_paragraph("MENTIONS LÉGALES", style='Heading 3')
    doc.add_paragraph(mentions_legales)
    
    doc.add_paragraph()
    
    echeance = doc.add_paragraph()
    echeance.add_run(f"Date d'échéance : {date_echeance}").italic = True


def generer_docx(devis: dict, nom_fichier: str = None) -> str:
    """
    Génère un document DOCX à partir d'un devis au format JSON.
    Retourne le chemin du fichier généré.
    """
    
    if nom_fichier is None:
        nom_fichier = f"{devis.get('numero_devis', 'devis')}.docx"
    
    chemin_complet = os.path.join(DOSSIER_SORTIE, nom_fichier)
    
    doc = Document()
    
    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Logo en haut
    ajouter_logo(doc)
    
    # Titre du document
    titre_doc = doc.add_paragraph()
    titre_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre_doc.add_run(f"DEVIS {devis.get('numero_devis', '')}")
    run.bold = True
    run.font.size = Pt(16)
    
    # Date
    date_par = doc.add_paragraph()
    date_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_par.add_run(f"Émis le {devis.get('date_emission', '')}")
    
    doc.add_paragraph()
    
    # Contenu
    ajouter_entete_dupsolaz(doc)
    ajouter_infos_emetteur(doc, devis.get('emetteur', {}))
    ajouter_infos_client(doc, devis.get('client', {}))
    ajouter_tableau_prestations(
        doc,
        devis.get('prestations', []),
        devis.get('total_ht', 0),
        devis.get('tgc', 6),
        devis.get('total_ttc', 0)
    )
    ajouter_zone_signature(doc)
    ajouter_mentions(
        doc,
        devis.get('conditions', []),
        devis.get('mentions_legales', ''),
        devis.get('date_echeance', '')
    )
    
    # Pied de page
    doc.add_paragraph()
    pied = doc.add_paragraph()
    pied.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pied.add_run("Dupsolaz Legacy — Merci de votre confiance.")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.save(chemin_complet)
    return chemin_complet


# ============================================================
# TEST
# ============================================================
if __name__ == "__main__":
    devis_test = {
        "numero_devis": "DEV-2026-07-001",
        "date_emission": "2026-07-09",
        "date_echeance": "2026-08-08",
        "emetteur": {
            "nom": "Dupsolaz Legacy",
            "ridet": "1 637 735.001",
            "forme_juridique": "Entreprise Individuelle",
            "adresse": "162 rue Edmond Mathey Mont-Dore",
            "email": "Dupsolaz0@gmail.com",
            "telephone": "+687 80 15 45"
        },
        "client": {
            "nom": "Jean Dupont",
            "adresse": "15 rue de Paris 75000 Paris",
            "email": "jean@dupont.fr",
            "telephone": "01 23 45 67 89"
        },
        "prestations": [
            {
                "libelle": "Prestation de secrétariat",
                "quantite": 10,
                "unite": "heure",
                "tarif_unitaire": 5000,
                "montant": 50000
            }
        ],
        "total_ht": 50000,
        "tgc": 6,
        "total_ttc": 53000,
        "conditions": [
            "Paiement à 30 jours à compter de la date d'émission du devis.",
            "Toute prestation commencée est due intégralement."
        ],
        "mentions_legales": "Dupsolaz Legacy — EI — RIDET 1 637 735.001 — Nouméa, Nouvelle-Calédonie."
    }
    
    print("Génération du document DOCX avec logo et signature...")
    chemin = generer_docx(devis_test)
    print(f"Document généré : {chemin}")